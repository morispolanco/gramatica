import os
import streamlit as st
import openai
import chardet
import docx
import PyPDF2
from io import BytesIO

openai.api_key = os.environ.get("OPENAI_API_KEY")

def leer_archivo(archivo_subido):
    extension = archivo_subido.name.split('.')[-1].lower()
    texto = ""

    if extension == 'txt':
        deteccion = chardet.detect(archivo_subido.read())
        codificacion = deteccion['encoding'] or 'latin-1'
        archivo_subido.seek(0)
        texto = archivo_subido.read().decode(codificacion)
    elif extension == 'docx':
        documento = docx.Document(archivo_subido)
        for p in documento.paragraphs:
            texto += p.text + "\n"
    elif extension == 'pdf':
        lector_pdf = PyPDF2.PdfFileReader(archivo_subido)
        for i in range(lector_pdf.getNumPages()):
            pagina = lector_pdf.getPage(i)
            texto += pagina.extractText()

    return texto

def dividir_texto(texto, max_tokens):
    palabras = texto.split()
    fragmentos = []
    fragmento_actual = []

    tokens_actual = 0
    for palabra in palabras:
        tokens_palabra = len(palabra) + 1  # Contar el espacio en blanco antes de la palabra
        if tokens_actual + tokens_palabra > max_tokens:
            fragmentos.append(' '.join(fragmento_actual))
            fragmento_actual = []
            tokens_actual = 0

        fragmento_actual.append(palabra)
        tokens_actual += tokens_palabra

    if fragmento_actual:
        fragmentos.append(' '.join(fragmento_actual))

    return fragmentos

def corregir_gramatica_y_estilo(texto_ensayo):
    max_tokens_por_fragmento = 3000
    fragmentos = dividir_texto(texto_ensayo, max_tokens_por_fragmento)
    correcciones = []

    for fragmento in fragmentos:
        prompt = f"Por favor, corrige la gramática y el estilo del siguiente ensayo:\n\n{fragmento}\n\nEnsayo corregido:"
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=150,
            n=1,
            stop=None,
            temperature=0.5,
        )
        correccion = response.choices[0].text.strip()
        correcciones.append(correccion)

    ensayo_corregido = ' '.join(correcciones)
    return ensayo_corregido

def guardar_ensayo_corregido(extension, ensayo_corregido):
    if extension == "txt":
        archivo_corregido = BytesIO()
        archivo_corregido.write(ensayo_corregido.encode("utf-8"))
        archivo_corregido.seek(0)
    elif extension == "docx":
        archivo_corregido = BytesIO()
        doc = docx.Document()
        doc.add_paragraph(ensayo_corregido)
        doc.save(archivo_corregido)
        archivo_corregido.seek(0)
    elif extension == "pdf":
        # (Agrega aquí la lógica de conversión de texto a PDF)
        # ...
        pass
    return archivo_corregido

st.title("Corrector de Gramática y Estilo")
archivo_subido = st.file_uploader("Sube tu ensayo", type=["txt", "docx", "pdf"])

if archivo_subido is not None:
    extension = archivo_subido.name.split(".")[-1].lower()
    texto_ensayo = leer_archivo(archivo_subido)

    if st.button("Corregir Gramática y Estilo", key="corregir_button"):
        ensayo_corregido = corregir_gramatica_y_estilo(texto_ensayo)
        archivo_corregido = guardar_ensayo_corregido(extension, ensayo_corregido)
        nombre_archivo_corregido = f"corregido_{archivo_subido.name}"

        if extension == "txt":
            mime_type = "text/plain"
        elif extension == "docx":
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif extension == "pdf":
            mime_type = "application/pdf"

        st.download_button("Descargar ensayo corregido", archivo_corregido.getvalue(), file_name=nombre_archivo_corregido, mime=mime_type)
